"""
Regenerate backend/docs/工單模板.docx with docxtpl placeholders (catalog keys).

Run from backend/: PYTHONPATH=. python scripts/build_work_order_template.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.domain.order_fields import DEFAULT_FIELD_ORDER, get_field_label

# Label (Traditional Chinese UI) -> catalog key for docxtpl
FIELD_ROWS: list[tuple[str, str]] = [
    (get_field_label(key), key) for key in DEFAULT_FIELD_ORDER
]

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "工單模板.docx"


def main() -> None:
    doc = Document()
    title = doc.add_paragraph("訂單工單")
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True

    table = doc.add_table(rows=len(FIELD_ROWS), cols=2)
    table.style = "Table Grid"
    for row_idx, (label, key) in enumerate(FIELD_ROWS):
        row = table.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = f"{{{{ {key} }}}}"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
